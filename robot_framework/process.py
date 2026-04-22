"""This module contains the main process of the robot."""
from sqlalchemy import create_engine, text
import pandas as pd
from docx import Document
import os
from .case_handler import CaseHandler
from .helper_functions import contact_lookup, check_case_folder, get_salary_case_id_through_metadata
from mbu_dev_shared_components.getorganized.objects import CaseDataJson
from .document_handler import DocumentHandler
from .journalize_process import journalize_file
from io import BytesIO
from docx.shared import RGBColor
from pathlib import Path
from OpenOrchestrator.orchestrator_connection.connection import OrchestratorConnection
from OpenOrchestrator.database.queues import QueueElement


def fetch_pending_employees(engine):
    query = text("""
        SELECT [Id],[Medarbejdernr],[Attesttype],[AttestModtagetDato]
            ,[RequestNumberServiceNow],[CPR],[Navn]
        FROM [RPA_BA].[attest].[AttestIndhentning]
        WHERE (LEN(Medarbejdernr) > 0
            AND Medarbejdernr <> '0'
            AND AttestModtaget = 1
            AND AttestStatus = 1
            AND (
                    (SendtTilPMappeQueue = 0 OR SendtTilPMappeQueue IS NULL)
                    OR
                    (SendtTilPMappeQueueDato < DATEADD(DAY, -2, GETDATE()))
            ))
            ORDER BY Id ASC;
            """)
    return pd.read_sql_query(query, con=engine)


def update_sql_information(engine, atteststatus, RequestNumberServiceNow, attesttype_value):
    query = text("""
        UPDATE [RPA_BA].[attest].[AttestIndhentning]
        SET [BehandletPMappeDato] = CURRENT_TIMESTAMP,
            [AttestStatus] = :atteststatus
        WHERE RequestNumberServiceNow = :request_number
        AND attesttype = (:attest_types)
        """)
    with engine.begin() as conn:
        conn.execute(query, {
            "atteststatus": atteststatus,
            "request_number": RequestNumberServiceNow,
            "attest_types": attesttype_value
        })


def udfyld_word_ark(tjenestenummer, fornavn, efternavn, attest_modtaget_dato, attesttype_value, att_id, temp_dir):
    if attesttype_value == 1:
        attest_type = "Privat straffeattest"
    elif attesttype_value == 2:
        attest_type = "Offentlig straffeattest"
    else:
        attest_type = "Børneattest"

    erstatninger = {
        "TJENESTENUMMER": tjenestenummer,
        "Tjenestenummer": tjenestenummer,
        "#SD-PERSON-FORNAVN#": fornavn,
        "#SD-PERSON-EFTERNAVN#": efternavn,
        "ATTESTMODTAGETDATO": str(attest_modtaget_dato),
        "TYPE-AF-ATTEST": attest_type,
    }

    doc = Document(Path(__file__).resolve().parent / 'AttestSkabelon.docx')

    def replace_in_runs(run):
        if not run.text:
            return
        txt = run.text.replace(" ", " ")
        for k, v in erstatninger.items():
            if k in txt:
                txt = txt.replace(k, str(v))
        run.text = txt
        run.font.color.rgb = RGBColor(0, 0, 0)

    for p in doc.paragraphs:
        for r in p.runs:
            replace_in_runs(r)

    def do_tables(tables):
        for t in tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            replace_in_runs(r)

    do_tables(doc.tables)
    for sec in doc.sections:
        do_tables(sec.header.tables)
        do_tables(sec.footer.tables)

    safe_date = attest_modtaget_dato.strftime("%d.%m.%Y")
    file_name = f"{attest_type} {safe_date} {att_id}.docx"
    file_path = temp_dir / file_name
    doc.save(file_path)
    return file_path


def find_personale_mappe(go_api_endpoint, go_api_username, go_api_password, cpr_dicts):
    try:
        case_handler = CaseHandler(go_api_endpoint, go_api_username, go_api_password)
        case_data_handler = CaseDataJson()

        cpr = next(iter(cpr_dicts.keys()))

        person_full_name, person_go_id = contact_lookup(case_handler, cpr)

        cases_info = check_case_folder(
            case_data_handler=case_data_handler,
            case_handler=case_handler,
            case_type="PER",
            person_full_name=person_full_name,
            person_go_id=person_go_id,
            ssn=cpr
        )

        if not cases_info:
            return None

        employee_folder_id = cases_info[0].get("CaseID")

        per_mappe_id = get_salary_case_id_through_metadata(
            case_handler=case_handler,
            employee_folder_id=employee_folder_id,
            case_title="Ansættelse og lønaftaler"
        )

        return per_mappe_id
    except Exception as e:
        cpr = next(iter(cpr_dicts.keys()), "UKENDT_CPR")
        print(f"Fejl ved forsøg på at finde per_mappe_id for CPR {cpr}: {e}")
        return None


def gem_fil_i_per_mappe(go_api_endpoint, go_api_username, go_api_password, per_mappe_id, fil_sti):
    try:
        document_handler = DocumentHandler(go_api_endpoint, go_api_username, go_api_password)

        with open(fil_sti, "rb") as f:
            file_bytes = f.read()

        filnavn = os.path.basename(fil_sti)
        salary_document_to_journalize_as_byte_stream = BytesIO(file_bytes)

        filename_with_extension = filnavn
        filename_without_extension = os.path.splitext(filnavn)[0]

        journalized_file_doc_id, status_message = journalize_file(
            document_category="Indgående",
            document_handler=document_handler,
            case_id=per_mappe_id,
            filename_with_extension=filename_with_extension,
            filename_without_extension=filename_without_extension,
            salary_document_to_journalize_as_byte_stream=salary_document_to_journalize_as_byte_stream
        )
        print(status_message)
        print(journalized_file_doc_id)
    except Exception as e:
        raise ValueError("Kunne ikke uploade pdf-fil til per_mappe", e)


# pylint: disable-next=unused-argument
def process(orchestrator_connection: OrchestratorConnection, queue_element: QueueElement | None = None) -> None:
    """Do the primary process of the robot."""
    orchestrator_connection.log_trace("Running process.")

    #Henter credentials til API bruger til GO:
    go_api_endpoint = orchestrator_connection.get_constant("GO_API_ENDPOINT")
    go_api_username = orchestrator_connection.get_credential("BA_GO_API").username
    go_api_password = orchestrator_connection.get_credential("BA_GO_API").password

    #Sætter stien til en temp mappe som notatet kan gemmes lokalt i
    TEMP_DIR = Path(r"C:\temp\attest_upload")
    TEMP_DIR.mkdir(parents=True, exist_ok=True)

    counter = 0
    #Forbindelse til databasen:
    engine = create_engine("mssql+pyodbc://faellessql/RPA_BA?driver=ODBC+Driver+17+for+SQL+Server&trusted_connection=yes")
    #Henter sager i kø og returnerer i dataframe:
    df = fetch_pending_employees(engine)

    if df.empty:
        orchestrator_connection.log_info("Ingen afventende medarbejdere at behandle.")
        return

    try:
        for rows in df.itertuples(index=False):
            try:
                tjenestenummer = str(rows.Medarbejdernr).strip()
                orchestrator_connection.log_trace(f"Behandler tjenestenummer: {tjenestenummer}")
                cpr = str(rows.CPR).strip()
                RequestNumberServiceNow = str(rows.RequestNumberServiceNow).strip()
                attesttype_value = int(rows.Attesttype)
                attest_modtaget_dato = rows.AttestModtagetDato
                name = str(rows.Navn).strip()
                att_id = str(rows.Id)

                fornavn, efternavn = name.split(" ", 1)
                file_path = udfyld_word_ark(tjenestenummer, fornavn, efternavn, attest_modtaget_dato, attesttype_value, att_id, TEMP_DIR)
                cpr_dict = {cpr: {"tjenestenummer": str(tjenestenummer), "navn": "", "stilling": ""}}
                personale_mappe = find_personale_mappe(go_api_endpoint, go_api_username, go_api_password, cpr_dict)

                if not personale_mappe:
                    file_path.unlink(missing_ok=True)
                    orchestrator_connection.log_info(f"Tjenestenummer {tjenestenummer} har ikke en personalemappe — springer over.")
                    continue

                gem_fil_i_per_mappe(go_api_endpoint, go_api_username, go_api_password, personale_mappe, file_path)
                update_sql_information(engine, 4, RequestNumberServiceNow, attesttype_value)
                file_path.unlink(missing_ok=True)
                counter += 1
                orchestrator_connection.log_trace(f"Færdig med nr. {counter}")

                if counter > 1:
                    break

            except Exception as e:
                orchestrator_connection.log_error(f"Fejl ved behandling af {tjenestenummer}: {e}")
                raise

    finally:
        for f in TEMP_DIR.glob("*.docx"):
            f.unlink(missing_ok=True)
